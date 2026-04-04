"""
DOCX 格式水印处理：嵌入/提取零宽字符。

仅负责文件格式操作（读写 DOCX 内部文本），
编解码逻辑由 office_wm.py 统一处理。
"""

from pathlib import Path
from typing import Optional

import docx
from loguru import logger


def embed_docx(
    input_path: Path, zwc_block: str, output_path: Path,
) -> tuple[bool, str]:
    """
    在 DOCX 文件中嵌入零宽字符块。

    找到第一个非空 run，在其文本前端插入 ZWC 块。
    修改 run.text 而非 paragraph.text 以保持格式。

    Returns:
        (success, message) 元组
    """
    try:
        document = docx.Document(str(input_path))
    except Exception as e:
        return False, f"Cannot open DOCX: {e}"

    # 查找第一个非空 run 并注入 ZWC
    inserted = False
    for para in document.paragraphs:
        for run in para.runs:
            if run.text and run.text.strip():
                run.text = zwc_block + run.text
                inserted = True
                break
        if inserted:
            break

    if not inserted:
        return False, "No text content found in DOCX"

    try:
        document.save(str(output_path))
    except Exception as e:
        return False, f"Cannot save DOCX: {e}"

    logger.debug(f"ZWC block inserted into DOCX: {input_path.name}")
    return True, "OK"


def extract_docx(file_path: Path) -> Optional[str]:
    """
    从 DOCX 文件提取所有文本内容。

    拼接所有 paragraph 的所有 run.text，
    返回完整文本供 zwc_decode 搜索水印标记。

    Returns:
        拼接后的全文文本，失败返回 None
    """
    try:
        document = docx.Document(str(file_path))
    except Exception as e:
        logger.warning(f"Cannot open DOCX for extraction: {e}")
        return None

    parts = []
    for para in document.paragraphs:
        for run in para.runs:
            if run.text:
                parts.append(run.text)
    return "".join(parts) if parts else None
